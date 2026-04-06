"""
nlp/text_extractor.py — Extract text from PDF, DOCX, and URLs
"""
import io
import re
from pathlib import Path
from typing import Optional
import requests
import pdfplumber
from docx import Document
from bs4 import BeautifulSoup
from backend.core.logging_config import get_logger

logger = get_logger(__name__)


class TextExtractor:
    """Unified text extractor supporting PDF, DOCX, TXT, and URL sources."""

    SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".doc", ".txt"}

    def extract(self, source: str | Path | bytes, source_type: str = "auto") -> str:
        """
        Extract plain text from a file path, bytes, or URL.

        Args:
            source: File path, raw bytes, or URL string.
            source_type: One of 'pdf', 'docx', 'txt', 'url', 'auto'.

        Returns:
            Cleaned text string.
        """
        if isinstance(source, bytes):
            return self._extract_from_bytes(source, source_type)

        source = str(source)

        if source_type == "url" or (source_type == "auto" and source.startswith("http")):
            return self._extract_from_url(source)

        path = Path(source)
        if not path.exists():
            raise FileNotFoundError(f"File not found: {source}")

        suffix = path.suffix.lower()
        if source_type == "auto":
            source_type = suffix.lstrip(".")

        return self._extract_by_type(path.read_bytes(), source_type)

    def _extract_from_bytes(self, data: bytes, source_type: str) -> str:
        """Extract text from raw bytes."""
        return self._extract_by_type(data, source_type)

    def _extract_by_type(self, data: bytes, source_type: str) -> str:
        source_type = source_type.lower().lstrip(".")

        # ── Magic Byte Verification ───────────────────────────
        if source_type == "pdf" and not data.startswith(b"%PDF-"):
            logger.warning("File has .pdf extension but missing %PDF- header. Falling back to text.")
            source_type = "txt"
        elif source_type in ("docx", "doc") and not data.startswith(b"PK\x03\x04"):
            # DOCX is a ZIP file starting with PK\x03\x04
            logger.warning(f"File has .{source_type} extension but missing ZIP header. Falling back to text.")
            source_type = "txt"

        extractors = {
            "pdf": self._from_pdf,
            "docx": self._from_docx,
            "doc": self._from_docx,
            "txt": lambda d: d.decode("utf-8", errors="ignore"),
        }
        extractor = extractors.get(source_type, lambda d: d.decode("utf-8", errors="ignore"))
        try:
            text = extractor(data)
        except Exception as exc:
            logger.error(f"Extraction failed for {source_type}, trying text fallback", error=str(exc))
            text = data.decode("utf-8", errors="ignore")

        return self._clean(text)

    def _from_pdf(self, data: bytes) -> str:
        """Extract text from PDF bytes using pdfplumber."""
        if not data:
            logger.error("PDF data is empty")
            return ""

        logger.debug("Starting PDF extraction", size=len(data), head=data[:50])
        text_parts = []
        try:
            with pdfplumber.open(io.BytesIO(data)) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text() or ""
                    text_parts.append(page_text)
        except Exception as exc:
            logger.error("PDF extraction failed", error=str(exc), data_head=data[:100])
            raise ValueError(f"PDF extraction error: {exc}")
        return "\n".join(text_parts)

    def _from_docx(self, data: bytes) -> str:
        """Extract text from DOCX bytes using python-docx."""
        try:
            doc = Document(io.BytesIO(data))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            paragraphs.append(cell.text.strip())
            return "\n".join(paragraphs)
        except Exception as exc:
            logger.error("DOCX extraction failed", error=str(exc))
            raise

    def _extract_from_url(self, url: str, timeout: int = 15) -> str:
        """Scrape text from a URL (job posting pages)."""
        try:
            headers = {
                "User-Agent": (
                    "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
                    "AppleWebKit/537.36 Chrome/120.0 Safari/537.36"
                )
            }
            resp = requests.get(url, headers=headers, timeout=timeout)
            resp.raise_for_status()
            soup = BeautifulSoup(resp.text, "lxml")

            # Remove script/style tags
            for tag in soup(["script", "style", "nav", "footer", "header"]):
                tag.decompose()

            return self._clean(soup.get_text(separator="\n"))
        except Exception as exc:
            logger.error("URL extraction failed", url=url, error=str(exc))
            raise

    @staticmethod
    def _clean(text: str) -> str:
        """Remove excessive whitespace and non-printable characters."""
        text = re.sub(r"\r\n|\r", "\n", text)
        text = re.sub(r"[ \t]+", " ", text)
        text = re.sub(r"\n{3,}", "\n\n", text)
        text = re.sub(r"[^\x20-\x7E\n]", " ", text)
        return text.strip()


# Singleton
text_extractor = TextExtractor()
